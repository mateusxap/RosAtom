import os
import random
from collections import OrderedDict
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm, Inches
from docx.oxml import OxmlElement
from faker import Faker
import matplotlib.pyplot as plt

# Инициализация Faker
locales = OrderedDict([
    ('en-US', 1),
    ('ru-RU', 2),
])
fake = Faker(locales)

# Путь к папке с изображениями
images_folder = 'natural_images'

# Создаем папку для сохранения документов
if not os.path.exists('docx'):
    os.makedirs('docx')

# Создаем папку для сохранения формул
if not os.path.exists('equations'):
    os.makedirs('equations')

num_documents = 10  # Количество документов для генерации

# Список цветов в формате HEX
COLORS = """000000 000080 00008B 0000CD 0000FF 006400 008000 008080 008B8B 00BFFF 00CED1 
00FA9A 00FF00 00FF7F 00FFFF 00FFFF 191970 1E90FF 20B2AA 228B22 2E8B57 2F4F4F 32CD32 3CB371 
40E0D0 4169E1 4682B4 483D8B 48D1CC 4B0082 556B2F 5F9EA0 6495ED 66CDAA 696969 6A5ACD 6B8E23 
708090 778899 7B68EE 7CFC00 7FFF00 7FFFD4 800000 800080 808000 808080 87CEEB 87CEFA 8A2BE2 
8B0000 8B008B 8B4513 8FBC8F 90EE90 9370D8 9400D3 98FB98 9932CC 9ACD32 A0522D A52A2A A9A9A9 
ADD8E6 ADFF2F AFEEEE B0C4DE B0E0E6 B22222 B8860B BA55D3 BC8F8F BDB76B C0C0C0 C71585 CD5C5C 
CD853F D2691E D2B48C D3D3D3 D87093 D8BFD8 DA70D6 DAA520 DC143C DCDCDC DDA0DD DEB887 E0FFFF 
E6E6FA E9967A EE82EE EEE8AA F08080 F0E68C F0F8FF F0FFF0 F0FFFF F4A460 F5DEB3 F5F5DC F5F5F5 
F5FFFA F8F8FF FA8072 FAEBD7 FAF0E6 FAFAD2 FDF5E6 FF0000 FF00FF FF00FF FF1493 FF4500 FF6347 
FF69B4 FF7F50 FF8C00 FFA07A FFA500 FFB6C1 FFC0CB FFD700 FFDAB9 FFDEAD FFE4B5 FFE4C4 FFE4E1 
FFEBCD FFEFD5 FFF0F5 FFF5EE FFF8DC FFFACD FFFAF0 FFFAFA FFFF00 FFFFE0 FFFFF0 FFFFFF"""
# Список LaTeX-формул
FORMULAS = [
    r"E = mc^2",
    r"\int_{a}^{b} f(x)\,dx",
    r"\frac{d}{dx}\left( e^x \right) = e^x",
    r"\sum_{n=1}^{\infty} \frac{1}{n^2} = \frac{\pi^2}{6}",
    r"\lim_{x \to 0} \frac{\sin x}{x} = 1",
    r"a^2 + b^2 = c^2",
    r"\nabla \cdot \mathbf{E} = \frac{\rho}{\varepsilon_0}",
    r"f(x) = \frac{1}{\sqrt{2\pi\sigma^2}} e^{-\frac{(x-\mu)^2}{2\sigma^2}}",
    r"i\hbar \frac{\partial}{\partial t}\Psi = \hat{H}\Psi",
    r"e^{i\theta} = \cos \theta + i \sin \theta",
    r"\frac{\partial^2 u}{\partial t^2} = c^2 \nabla^2 u",
    r"\alpha + \beta = \gamma",
    r"\sqrt{a^2 + b^2 + c^2}",
    r"\frac{1}{1 + e^{-x}}",
    r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}",
    r"\mathbf{F} = m\mathbf{a}",
    r"PV = nRT",
    r"\frac{\partial u}{\partial t} + \mathbf{v} \cdot \nabla u = D \nabla^2 u",
    r"\sigma = \frac{F}{A}",
    r"\tau = r \times F"
]


def generate_equation_image(equation_str, output_dir='equations'):
    """
    Генерирует изображение формулы из LaTeX-строки и сохраняет его.

    :param equation_str: Строка LaTeX-формулы.
    :param output_dir: Директория для сохранения изображений формул.
    :return: Путь к сохранённому изображению формулы.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Генерируем уникальное имя файла
    filename = f"equation_{random.randint(1000, 9999)}.png"
    filepath = os.path.join(output_dir, filename)

    # Создаём изображение формулы
    plt.figure(figsize=(3, 1))
    plt.text(0.5, 0.5, f"${equation_str}$", fontsize=20, ha='center', va='center')
    plt.axis('off')
    plt.savefig(filepath, bbox_inches='tight', pad_inches=0.1)
    plt.close()

    return filepath


def add_equation_to_docx(doc, equation_str, caption=True):
    """
    Генерирует изображение формулы и добавляет его в документ с меткой '~' белым цветом перед формулой.
    """
    equation_image_path = generate_equation_image(equation_str)
    try:
        if caption:
            # Создаем метку с символом '~' перед формулой
            label_paragraph = doc.add_paragraph("~")
            label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = label_paragraph.runs[0]
            run.font.size = Pt(1)  # Настройте размер шрифта при необходимости
            run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет шрифта

            # Добавляем свойство 'keep_with_next' к параграфу с меткой
            label_paragraph.paragraph_format.keep_with_next = True

        # Добавляем изображение формулы
        doc.add_picture(equation_image_path, width=Inches(4))

        # Получаем параграф с изображением и устанавливаем 'keep_together'
        image_paragraph = doc.paragraphs[-1]
        image_paragraph.paragraph_format.keep_together = True

    except Exception as e:
        print(f"Ошибка при добавлении формулы: {e}")


def add_footnote(paragraph, footnote_text, footnote_num, footnotes):
    """
    Добавляет сноску в абзац и сохраняет её текст.

    :param paragraph: Объект абзаца.
    :param footnote_text: Текст сноски.
    :param footnote_num: Номер сноски.
    :param footnotes: Список сносок.
    """
    footnote_mark = paragraph.add_run(f'[{footnote_num}]')
    footnote_mark.font.superscript = True
    footnotes.append((footnote_num, footnote_text))


def add_footnotes_section(document, footnotes):
    """
    Добавляет раздел с примечаниями (сносками) в конец документа.

    :param document: Объект документа Document.
    :param footnotes: Список сносок.
    """
    if footnotes:
        document.add_page_break()
        footnote_heading = document.add_paragraph('Примечания')
        footnote_heading.style = 'Heading 1'
        for num, text in footnotes:
            footnote_paragraph = document.add_paragraph()
            footnote_run = footnote_paragraph.add_run(f'[{num}] {text}')
            footnote_run.font.size = Pt(8)


colors_list = COLORS.split()


def set_table_borders(table, include_borders=True):
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = tbl.new_tblPr()

    if include_borders:
        tblBorders = OxmlElement('w:tblBorders')

        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
    else:
        # Удаляем границы, если они есть
        for element in tblPr.xpath('w:tblBorders'):
            tblPr.remove(element)


for doc_num in range(num_documents):

    document = Document()

    # Инициализация списка сносок
    footnotes = []
    footnote_num = 1  # Начальный номер сноски

    # Случайное количество разделов в документе
    num_sections = random.randint(1, 5)
    for section_num in range(num_sections):
        if section_num != 0:
            document.add_section(WD_SECTION.NEW_PAGE)

        # Случайное изменение ориентации страницы
        if random.randint(0, 1):
            current_section = document.sections[-1]
            new_width, new_height = current_section.page_height, current_section.page_width
            current_section.orientation = WD_ORIENT.LANDSCAPE
            current_section.page_width = new_width
            current_section.page_height = new_height

        # Случайно выбираем, использовать ли несколько колонок
        use_columns = random.choice([True, False])
        if use_columns:
            columns = random.choice([2, 3])
            sectPr = document.sections[-1]._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), str(columns))
        else:
            # Устанавливаем одну колонку
            sectPr = document.sections[-1]._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), '1')

        # Добавляем верхний колонтитул
        header = document.sections[-1].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = fake.sentence(nb_words=random.randint(5, 10))

        # Добавляем нижний колонтитул
        footer = document.sections[-1].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = fake.sentence(nb_words=random.randint(5, 10))

        # Добавляем заголовок
        level = random.randint(0, 4)
        heading_text = fake.sentence(nb_words=random.randint(3, 7))
        heading = document.add_heading(heading_text, level=level)
        run = heading.runs[0]
        run.font.bold = random.choice([True, False])
        run.font.italic = random.choice([True, False])
        heading_size = random.randint(14, 24)
        run.font.size = Pt(heading_size)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if random.choice([True, False]) else WD_ALIGN_PARAGRAPH.LEFT

        # Добавляем абзац текста с возможными сносками
        paragraph = document.add_paragraph(fake.text(max_nb_chars=random.randint(500, 1000)))
        # Выбираем размер шрифта для основного текста (как минимум на 2 пункта меньше заголовка)
        max_text_size = heading_size - 2
        font_size = Pt(random.randint(8, heading_size - 2))
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Cm(1) if random.choice([True, False]) else None
        paragraph_format.alignment = random.choice([
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY
        ])

        # Выбираем единый размер шрифта для всего абзаца
        font_size = Pt(random.randint(8, heading_size - 2))

        # Разбиваем текст на предложения
        sentences = paragraph.text.split('. ')
        paragraph.text = ''  # Очищаем текст абзаца для повторного заполнения
        for sentence in sentences:
            if sentence.strip() == '':
                continue
            run = paragraph.add_run(sentence + '. ')
            run.font.size = font_size  # Применяем единый размер шрифта
            run.font.name = 'Times New Roman'

            # Случайно добавляем сноску
            if random.choice([True, False, False]):  # Увеличиваем вероятность добавления сносок
                footnote_text = fake.sentence(nb_words=5)
                add_footnote(paragraph, footnote_text, footnote_num, footnotes)
                footnote_num += 1

        # Добавляем подпись к таблице
        caption_text = f"Таблица {random.randint(1, 100)} — {fake.sentence(nb_words=random.randint(3, 7))}"
        caption_paragraph = document.add_paragraph(caption_text)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.runs[0]
        run.font.size = Pt(random.randint(10, 12))

        # Добавляем таблицу
        table = document.add_table(
            rows=random.randint(2, 5),
            cols=random.randint(2, 5)
        )
        table.alignment = random.choice([
            WD_TABLE_ALIGNMENT.LEFT,
            WD_TABLE_ALIGNMENT.CENTER,
            WD_TABLE_ALIGNMENT.RIGHT
        ])

        # Решаем, делать ли таблицу цветной
        make_colorful = random.choice([True, False])

        # Решаем, будут ли границы таблицы
        include_borders = random.choice([True, False])
        set_table_borders(table, include_borders=include_borders)

        if make_colorful:
            # Выбираем случайные цвета для строк
            color_row_1 = random.choice(colors_list)
            color_row_2 = random.choice(colors_list)

            # Выбираем цвет текста: белый или черный
            font_color = RGBColor(255, 255, 255) if random.choice([True, False]) else RGBColor(0, 0, 0)

            for idx_row, row in enumerate(table.rows):
                # Выбираем цвет для текущей строки
                color = color_row_1 if idx_row % 2 == 0 else color_row_2

                for cell in row.cells:
                    # Заполняем ячейку текстом
                    cell.text = fake.word()

                    # Устанавливаем цвет заливки ячейки
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), color)
                    cell._tc.get_or_add_tcPr().append(shading_elm)

                    # Устанавливаем цвет и размер шрифта
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = font_color
                            run.font.size = Pt(random.randint(8, 14))

                    # Устанавливаем выравнивание содержимого ячейки
                    alignment = random.choice([
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.RIGHT,
                        WD_ALIGN_PARAGRAPH.JUSTIFY
                    ])
                    cell.paragraphs[0].alignment = alignment
        else:
            # Если таблица не цветная, заполняем ячейки стандартно
            for row in table.rows:
                for cell in row.cells:
                    cell.text = fake.word()

                    # Устанавливаем размер шрифта
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(random.randint(8, 14))

                    # Устанавливаем выравнивание содержимого ячейки
                    alignment = random.choice([
                        WD_ALIGN_PARAGRAPH.LEFT,
                        WD_ALIGN_PARAGRAPH.CENTER,
                        WD_ALIGN_PARAGRAPH.RIGHT,
                        WD_ALIGN_PARAGRAPH.JUSTIFY
                    ])
                    cell.paragraphs[0].alignment = alignment

        # Добавляем изображение с подписью и меткой '&' перед изображением
        image_files = os.listdir(images_folder) if os.path.exists(images_folder) else []
        if image_files:
            image_path = os.path.join(images_folder, random.choice(image_files))
            try:
                # Добавляем метку с символом '&' перед изображением
                label_paragraph = document.add_paragraph("&")
                label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = label_paragraph.runs[0]
                run.font.size = Pt(1)  # Настройте размер шрифта при необходимости
                run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет шрифта

                # Устанавливаем свойство 'keep_with_next' для метки
                label_paragraph.paragraph_format.keep_with_next = True

                # Добавляем изображение
                document.add_picture(image_path, width=Inches(4))

                # Получаем параграф с изображением и устанавливаем 'keep_together'
                image_paragraph = document.paragraphs[-1]
                image_paragraph.paragraph_format.keep_together = True

            except Exception as e:
                print(f"Ошибка при добавлении изображения: {e}")

            caption_text = f"Рисунок {random.randint(1, 100)} — {fake.sentence(nb_words=random.randint(3, 7))}"
            caption_paragraph = document.add_paragraph(caption_text)
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_paragraph.runs[0]
            run.font.size = Pt(random.randint(10, 12))

        # Добавляем нумерованный список
        for _ in range(random.randint(3, 7)):
            list_item = fake.sentence(nb_words=random.randint(5, 15))
            paragraph = document.add_paragraph(list_item, style='List Number')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Cm(1)
            paragraph_format.line_spacing = random.uniform(1.0, 1.5)
            run = paragraph.runs[0]
            run.font.size = Pt(random.randint(8, 16))
            run.font.name = 'Times New Roman'

        # Добавляем маркированный список
        for _ in range(random.randint(3, 7)):
            list_item = fake.sentence(nb_words=random.randint(5, 15))
            paragraph = document.add_paragraph(list_item, style='List Bullet')
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Cm(1)
            paragraph_format.line_spacing = random.uniform(1.0, 1.5)
            run = paragraph.runs[0]
            run.font.size = Pt(random.randint(8, 16))
            run.font.name = 'Times New Roman'

        # Добавляем абзац текста с возможными сносками
        paragraph = document.add_paragraph(fake.text(max_nb_chars=random.randint(500, 1000)))
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Cm(1) if random.choice([True, False]) else None
        paragraph_format.alignment = random.choice([
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.RIGHT,
            WD_ALIGN_PARAGRAPH.JUSTIFY
        ])

        # Выбираем единый размер шрифта для всего абзаца
        font_size = Pt(random.randint(8, heading_size - 2))

        # Разбиваем текст на предложения
        sentences = paragraph.text.split('. ')
        paragraph.text = ''  # Очищаем текст абзаца для повторного заполнения
        for sentence in sentences:
            if sentence.strip() == '':
                continue
            run = paragraph.add_run(sentence + '. ')
            run.font.size = font_size  # Применяем единый размер шрифта
            run.font.name = 'Times New Roman'

            # Случайно добавляем сноску
            if random.choice([True, False, False]):  # Увеличиваем вероятность добавления сносок
                footnote_text = fake.sentence(nb_words=5)
                add_footnote(paragraph, footnote_text, footnote_num, footnotes)
                footnote_num += 1

        # Добавляем формулу с возможной подписью
        if random.choice([True, False]):
            equation_str = random.choice(FORMULAS)
            add_equation_to_docx(document, equation_str, caption=True)

    # Добавляем сноски в конец документа
    add_footnotes_section(document, footnotes)

    # Сохраняем документ в папку 'docx'
    document.save(f'docx/document_{doc_num}.docx')
    print(f"Документ {doc_num} успешно сгенерирован.")