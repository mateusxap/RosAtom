import os
import random
from collections import OrderedDict
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt, Cm, Inches, Emu
from docx.oxml import OxmlElement
from faker import Faker
import matplotlib.pyplot as plt
from PIL import Image
import numpy as np

# Инициализация Faker
locales = OrderedDict([
    ('en-US', 1),
    ('ru-RU', 2),
])
fake = Faker(locales)

# Путь к папке с изображениями
images_folder = 'natural_images'

# Создаем папки для сохранения различных ресурсов
os.makedirs('docx', exist_ok=True)
os.makedirs('equations', exist_ok=True)
os.makedirs('plots', exist_ok=True)

num_documents = 3  # Количество документов для генерации

# Список цветов в формате HEX
COLORS = """000000 000080 00008B 0000CD 0000FF 006400 008000 008080 008B8B 00BFFF 00CED1 
00FA9A 00FF00 00FF7F 00FFFF 00FFFF 191970 1E90FF 20B2AA 228B22 2E8B57 2F4F4F 32CD32 3CB371 
40E0D0 4169E1 4682B4 483D8B 48D1CC 4B0082 556B2F 5F9EA0 6495ED 66CDAA 696969 6A5ACD 6B8E23 
708090 778899 7B68EE 7CFC00 7FFF00 7FFFD4 800000 800080 808000 808080 87CEEB 87CEFA 8A2BE2 
8B0000 8B008B 8B4513 8FBC8F 90EE90 9370D8 9400D3 98FB98 9932CC 9ACD32 A0522D A52A2A A9A9A9 
ADD8E6 ADFF2F AFEEEE B0C4DE B0E0E6 B22222 B8860B BA55D3 BC8F8F BDB76B C0C0C0 C71585 CD5C5C 
CD853F D2691E D2B48C D3D3D3 D87093 D8BFD8 DA70D6 DAA520 DC143C DCDCDC DDA0DD DEB887 E0FFFF 
E6E6FA E9967A EE82EE EEE8AA F08080 F0E68C F0F8FF F0FFF0 F0FFFF F4A460 F5DEB3 F5F5DC F5F5F5 
F5FFFA F8F8FF FA8072 FAEBD7 FAF0E6 FAFAD2 FDF5E6 FF0000 FF00FF FF00FF FF1493 FF4500 FF6347 
FF69B4 FF7F50 FF8C00 FFA07A FFA500 FFB6C1 FFC0CB FFD700 FFDAB9 FFDEAD FFE4B5 FFE4C4 FFE4E1 
FFEB3B FFEBCD FFEFD5 FFF0F5 FFF5EE FFF8DC FFFACD FFFAF0 FFFAFA FFFF00 FFFFE0 FFFFF0 FFFFFF"""

# Список LaTeX-формул
FORMULAS = [
    r"E = mc^2",
    r"\int_{a}^{b} f(x)\,dx",
    r"\frac{d}{dx}\left( e^x \right) = e^x",
    r"\sum_{n=1}^{\infty} \frac{1}{n^2} = \frac{\pi^2}{6}",
    r"\lim_{x \to 0} \frac{\sin x}{x} = 1",
    r"a^2 + b^2 = c^2",
    r"\nabla \cdot \mathbf{E} = \frac{\rho}{\varepsilon_0}",
    r"f(x) = \frac{1}{\sqrt{2\pi\sigma^2}} e^{-\frac{(x-\mu)^2}{2\sigma^2}}",
    r"i\hbar \frac{\partial}{\partial t}\Psi = \hat{H}\Psi",
    r"e^{i\theta} = \cos \theta + i \sin \theta",
    r"\frac{\partial^2 u}{\partial t^2} = c^2 \nabla^2 u",
    r"\alpha + \beta = \gamma",
    r"\sqrt{a^2 + b^2 + c^2}",
    r"\frac{1}{1 + e^{-x}}",
    r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}",
    r"\mathbf{F} = m\mathbf{a}",
    r"PV = nRT",
    r"\frac{\partial u}{\partial t} + \mathbf{v} \cdot \nabla u = D \nabla^2 u",
    r"\sigma = \frac{F}{A}",
    r"\tau = r \times F"
]

def generate_equation_image(equation_str, output_dir='equations'):
    """
    Генерирует изображение формулы из LaTeX-строки и сохраняет его.

    :param equation_str: Строка LaTeX-формулы.
    :param output_dir: Директория для сохранения изображений формул.
    :return: Путь к сохранённому изображению формулы.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Генерируем уникальное имя файла
    filename = f"equation_{random.randint(1000, 9999)}.png"
    filepath = os.path.join(output_dir, filename)

    # Создаём изображение формулы
    plt.figure(figsize=(3, 1))
    plt.text(0.5, 0.5, f"${equation_str}$", fontsize=20, ha='center', va='center')
    plt.axis('off')
    plt.savefig(filepath, bbox_inches='tight', pad_inches=0.1)
    plt.close()

    return filepath

def add_image_to_document(document, image_path, max_height_px=500, base_font_size=12):
    """
    Добавляет изображение в документ с ограничением по высоте.

    :param document: Объект документа Document.
    :param image_path: Путь к изображению.
    :param max_height_px: Максимальная высота изображения в пикселях.
    :param base_font_size: Базовый размер шрифта для метки.
    :return: Параграф с изображением.
    """
    # Открываем изображение
    image = Image.open(image_path)
    width_px, height_px = image.size

    # Ограничиваем высоту изображения
    if height_px > max_height_px:
        scaling_factor = max_height_px / height_px
        width_px = int(width_px * scaling_factor)
        height_px = max_height_px

    # Переводим размеры из пикселей в EMU
    width_emu = Emu(width_px / 96 * Inches(1).emu)
    height_emu = Emu(height_px / 96 * Inches(1).emu)

    # Создаем параграф и добавляем метку и изображение
    paragraph = document.add_paragraph()
    run = paragraph.add_run()

    # Добавляем метку '&'
    run.add_text("&")
    run.font.size = Pt(base_font_size)
    run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет

    # Добавляем изображение
    run.add_picture(image_path, width=width_emu, height=height_emu)

    # Устанавливаем выравнивание параграфа
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Устанавливаем свойства форматирования параграфа
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = True

    return paragraph

def add_equation_to_docx(doc, equation_str, size_img, caption=True ):
    equation_image_path = generate_equation_image(equation_str)
    try:
        if caption:
            # Создаем параграф и добавляем метку и изображение
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()

            # Добавляем метку '~'
            run.add_text("~")
            run.font.size = Pt(1)
            run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет

            # Определяем случайный размер

            img = Image.open(equation_image_path)
            width, height = img.size

            if size_img == "normal":
                width *= 0.75
                height *= 0.75
            if size_img == "small":
                width *= 0.55
                height *= 0.55
            elif size_img == "smallest":
                width *= 0.35
                height *= 0.35

            # Добавляем изображение формулы с выбранным размером
            run.add_picture(equation_image_path, width=Inches(width/96), height=Inches(height/96)) # Перевод из px в inches



            # Устанавливаем выравнивание параграфа
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Устанавливаем свойства форматирования параграфа
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = True

    except Exception as e:
        print(f"Ошибка при добавлении формулы: {e}")


def add_footnote(paragraph, footnote_text, footnote_num, footnotes, base_font_size=12):
    """
    Добавляет сноску в абзац и сохраняет её текст.

    :param paragraph: Объект абзаца.
    :param footnote_text: Текст сноски.
    :param footnote_num: Номер сноски.
    :param footnotes: Список сносок.
    :param base_font_size: Базовый размер шрифта для сноски.
    """
    footnote_mark = paragraph.add_run(f' [{footnote_num}] ')
    footnote_mark.font.size = Pt(base_font_size)
    footnotes.append((footnote_num, footnote_text))  

def add_footnotes_section(document, footnotes, base_font_size=12):
    """
    Добавляет раздел с примечаниями (сносками) в конец документа.

    :param document: Объект документа Document.
    :param footnotes: Список сносок.
    :param base_font_size: Базовый размер шрифта для сносок.
    """
    if footnotes:
        document.add_page_break()
        footnote_heading = document.add_paragraph('Примечания')
        footnote_heading.style = 'Heading 1'
        footnote_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for num, text in footnotes:
            footnote_paragraph = document.add_paragraph()
            footnote_run = footnote_paragraph.add_run(f'[{num}] {text}')
            footnote_run.font.size = Pt(base_font_size)
            footnote_paragraph.paragraph_format.keep_with_next = True

colors_list = COLORS.split()

def set_table_borders(table, borders=None):
    """
    Устанавливает границы таблицы. Если borders пустой или None, границы не устанавливаются.

    :param table: Объект таблицы.
    :param borders: Список границ для установки (например, ['top', 'bottom', 'left', 'right', 'insideH', 'insideV']).
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    if tblPr is None:
        tblPr = tbl.new_tblPr()

    # Удаляем существующие границы
    for element in tblPr.xpath('w:tblBorders'):
        tblPr.remove(element)

    # Создаем элемент tblBorders
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        if borders and border_name in borders:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
        else:
            border.set(qn('w:val'), 'nil')  # Отсутствие границы
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_table_keep_together(table):
    """
    Устанавливает свойства для предотвращения разрыва таблицы на разных страницах или колонках.
    """
    for row in table.rows:
        # Устанавливаем свойство 'cantSplit' для строки
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                # Устанавливаем свойства для параграфов внутри ячеек
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = True

def generate_random_plot(output_dir='plots'):
    """
    Генерирует случайный график и сохраняет его как изображение.
    """
    os.makedirs(output_dir, exist_ok=True)

    # Генерируем уникальное имя файла
    filename = f"plot_{random.randint(1000, 9999)}.png"
    filepath = os.path.join(output_dir, filename)

    # Определяем диапазон x
    x = np.linspace(-5, 5, 100)  # 100 точек от -5 до 5

    # Список доступных функций
    functions = [
        np.sin, np.cos, np.tan, np.exp, lambda x: x ** 2, lambda x: x ** 3
    ]

    # Выбираем случайную функцию
    f = random.choice(functions)

    # Вычисляем y
    y = f(x)

    # Создаём график
    plt.figure()
    plt.plot(x, y)
    plt.title(f"График функции: {fake.word()}")
    plt.xlabel("Ось X")
    plt.ylabel("Ось Y")
    plt.grid(True)
    plt.savefig(filepath)
    plt.close()

    return filepath

def add_plot_to_docx(doc, base_font_size=12):
    """
    Добавляет график в документ.

    :param doc: Объект документа Document.
    :param base_font_size: Базовый размер шрифта для метки.
    """
    plot_image_path = generate_random_plot()
    try:
        # Создаем параграф и добавляем метку и изображение
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # Добавляем метку '$'
        run.add_text("$")
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет

        # Добавляем изображение графика
        run.add_picture(plot_image_path)

        # Устанавливаем выравнивание параграфа
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Устанавливаем свойства форматирования параграфа
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.keep_with_next = True

        # Добавляем подпись к графику
        caption_type = random.choice(["Рис.", "Рисунок", "Рис. №", "Рисунок №"])
        caption_text = f"{caption_type} {random.randint(1, 100)} — {fake.sentence(nb_words=random.randint(3, 7))}"
        caption_paragraph = doc.add_paragraph(caption_text)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_paragraph.runs[0]
        run.font.size = Pt(base_font_size)
        caption_paragraph.paragraph_format.keep_with_next = True

    except Exception as e:
        print(f"Ошибка при добавлении графика: {e}")

def set_numbering_font_size(document, base_font_size=12):
    """
    Устанавливает размер шрифта для номеров и маркеров всех списков в документе.

    :param document: Объект документа Document.
    :param base_font_size: Базовый размер шрифта.
    """
    numbering_part = document.part.numbering_part
    numbering_elem = numbering_part.element

    for abstractNum in numbering_elem.findall(qn('w:abstractNum')):
        for lvl in abstractNum.findall(qn('w:lvl')):
            rPr = lvl.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                lvl.append(rPr)
            sz = rPr.find(qn('w:sz'))
            if sz is None:
                sz = OxmlElement('w:sz')
                rPr.append(sz)
            sz_val = str(int(base_font_size * 2))  # Размер шрифта в половинных пунктах
            sz.set(qn('w:val'), sz_val)

for doc_num in range(num_documents):

    document = Document()

    # Определение базового размера шрифта и размера заголовка
    base_font_size = random.randint(10, 12)  # Вы можете изменить диапазон по необходимости
    heading_size = random.randint(base_font_size + 2, base_font_size + 6)  # Минимум на 2 больше

    footnotes = []
    footnote_num = 1  # Начальный номер сноски

    doc_line_spacing = random.uniform(1.0, 1.2)

    # Случайное количество разделов в документе
    num_sections = random.randint(3, 7)
    for section_num in range(num_sections):
        if section_num != 0:
            document.add_section(WD_SECTION.NEW_PAGE)

        current_section = document.sections[-1]
        
        # Случайное изменение ориентации страницы
        if random.randint(0, 1):
            new_width, new_height = current_section.page_height, current_section.page_width
            current_section.orientation = WD_ORIENT.LANDSCAPE
            current_section.page_width = new_width
            current_section.page_height = new_height

        # Случайно выбираем, использовать ли несколько колонок
        use_columns = random.choice([True, False])
        if use_columns:
            columns = random.choice([2, 3])
            sectPr = document.sections[-1]._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), str(columns))
        else:
            # Устанавливаем одну колонку
            sectPr = document.sections[-1]._sectPr
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'), '1')

        # Случайно решаем, добавлять ли верхний колонтитул
        if random.choice([True, False, False]):
            # Добавляем верхний колонтитул
            header = current_section.header
            header_paragraph = header.paragraphs[0]
            header_paragraph.text = fake.sentence(nb_words=random.randint(1, 6))
            # Устанавливаем размер шрифта для колонтитула
            for run in header_paragraph.runs:
                run.font.size = Pt(base_font_size)

        # Случайно решаем, добавлять ли нижний колонтитул
        if random.choice([True, False, False]):
            # Добавляем нижний колонтитул
            footer = current_section.footer
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = fake.sentence(nb_words=random.randint(1, 6))
            # Устанавливаем размер шрифта для колонтитула
            for run in footer_paragraph.runs:
                run.font.size = Pt(base_font_size)

        # Добавляем заголовок
        level = random.randint(0, 4)
        heading_text = fake.sentence(nb_words=random.randint(3, 7))
        heading = document.add_heading(heading_text, level=level)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        choise_italic = random.choice([True, False])
        run.font.italic = choise_italic
        run.font.bold = random.choice([True, False]) or (not choise_italic)
        run.font.size = Pt(heading_size)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if random.choice([True, False]) else WD_ALIGN_PARAGRAPH.LEFT

        # Определяем функции для добавления различных элементов
        def add_text_paragraph():
            # Инициализация списка сносок
            global footnotes
            global footnote_num  # Начальный номер сноски
            # Добавляем абзац текста с возможными сносками
            paragraph = document.add_paragraph(fake.text(max_nb_chars=random.randint(1000, 1500)))
            paragraph_format = paragraph.paragraph_format
            if random.choice([True, False]):
                paragraph_format.first_line_indent = Cm(1)
            paragraph_format.line_spacing = doc_line_spacing
            paragraph_format.alignment = random.choice([
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT,
                WD_ALIGN_PARAGRAPH.JUSTIFY
            ])

            # Устанавливаем единый размер шрифта для всего абзаца
            for run in paragraph.runs:
                run.font.size = Pt(base_font_size)
                run.font.name = 'Times New Roman'

                # Случайно добавляем сноску
                if random.choice([False, False, True, False, False]):
                    footnote_text = fake.sentence(nb_words=5)
                    add_footnote(paragraph, footnote_text, footnote_num, footnotes, base_font_size=base_font_size)
                    footnote_num += 1


        def add_table():
            # Рисуем таблицы только в таком случае, если нет колонок
            if use_columns:
                return
            table_sign_up = random.choice([True, False])
            if table_sign_up:
                # Добавляем подпись к таблице
                caption_type = random.choice(["Табл.", "Таблица", "Табл. №", "Таблица №",])
                caption_text = f"{caption_type} {random.randint(1, 100)} — {fake.sentence(nb_words=random.randint(3, 7))}"
                caption_paragraph = document.add_paragraph(caption_text)
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption_paragraph.runs[0]
                run.font.size = Pt(base_font_size)
                caption_paragraph.paragraph_format.keep_with_next = True

            # Добавляем таблицу
            table = document.add_table(
                rows=random.randint(2, 5),
                cols=random.randint(2, 5)
            )
            table.alignment = random.choice([
                WD_TABLE_ALIGNMENT.LEFT,
                WD_TABLE_ALIGNMENT.CENTER,
                WD_TABLE_ALIGNMENT.RIGHT
            ])

            # Решаем, какой тип таблицы создать
            table_type = random.choice(['colorful_no_grid', 'ordinary_with_grid'])

            # Устанавливаем общие параметры форматирования для ячеек
            cell_alignment = random.choice([
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.RIGHT,
                WD_ALIGN_PARAGRAPH.JUSTIFY
            ])

            if table_type == 'colorful_no_grid':
                # Таблица цветная без сетки
                set_table_borders(table, borders=[])  # Убираем все границы

                # Решаем, будет ли первая строка жирной
                first_row_bold = random.random() < 0.5  # 50% вероятность

                # Выбираем случайные цвета для строк
                color_row_1 = random.choice(colors_list)
                color_row_2 = random.choice(colors_list)

                # Выбираем цвет текста: белый или черный
                font_color = RGBColor(255, 255, 255) if random.choice([True, False]) else RGBColor(0, 0, 0)

                for idx_row, row in enumerate(table.rows):
                    # Выбираем цвет для текущей строки
                    color = color_row_1 if idx_row % 2 == 0 else color_row_2

                    # Решаем, заполнять ли строку словами или числами
                    if idx_row == 0:
                        fill_with_words = True  # Первая строка всегда слова
                    else:
                        fill_with_words = random.choice([True, False])  # 50% вероятность

                    for idx_cell, cell in enumerate(row.cells):
                        # Заполнение ячеек
                        if fill_with_words:
                            cell.text = fake.word()
                        else:
                            cell.text = str(random.randint(1, 100))

                        # Устанавливаем цвет заливки ячейки
                        shading_elm = OxmlElement('w:shd')
                        shading_elm.set(qn('w:fill'), color)
                        cell._tc.get_or_add_tcPr().append(shading_elm)

                        # Устанавливаем цвет, размер шрифта и выравнивание
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = cell_alignment
                            for run in paragraph.runs:
                                run.font.color.rgb = font_color
                                run.font.size = Pt(base_font_size)

                    # Применяем жирный шрифт ко всей первой строке, если нужно
                    if idx_row == 0 and first_row_bold:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

            else:
                # Обычная таблица с сеткой
                set_table_borders(table, borders=['top', 'left', 'bottom', 'right', 'insideH', 'insideV'])

                # Решаем, будет ли первая строка жирной
                first_row_bold = random.random() < 0.5  # 50% вероятность

                # Для обычных таблиц шрифт всегда черный
                font_color = RGBColor(0, 0, 0)

                for idx_row, row in enumerate(table.rows):
                    # Решаем, заполнять ли всю строку словами или числами
                    if idx_row == 0:
                        fill_with_words = True  # Первая строка всегда слова
                    else:
                        fill_with_words = random.choice([True, False])  # 50% вероятность

                    for idx_cell, cell in enumerate(row.cells):
                        # Заполнение ячеек
                        if fill_with_words:
                            cell.text = fake.word()
                        else:
                            cell.text = str(random.randint(1, 100))

                        # Устанавливаем размер шрифта и выравнивание
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = cell_alignment
                            for run in paragraph.runs:
                                run.font.size = Pt(base_font_size)
                                run.font.name = 'Times New Roman'
                                run.font.color.rgb = font_color

                    # Применяем жирный шрифт ко всей первой строке, если нужно
                    if idx_row == 0 and first_row_bold:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

            # Устанавливаем свойства для сохранения целостности таблицы
            set_table_keep_together(table)

            if not table_sign_up:
                # Добавляем подпись к таблице
                caption_text = f"Таблица {random.randint(1, 100)} — {fake.sentence(nb_words=random.randint(3, 7))}"
                caption_paragraph = document.add_paragraph(caption_text)
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption_paragraph.runs[0]
                run.font.size = Pt(base_font_size)
                caption_paragraph.paragraph_format.keep_with_next = True

        def add_image_or_graph():
            if use_columns:
                return
            add_graph = random.choice([True, False])
            if add_graph:
                # Добавляем график
                add_plot_to_docx(document, base_font_size=base_font_size)


            add_image = random.choice([True, False])
            if add_image:
                # Добавляем изображение
                image_files = os.listdir(images_folder) if os.path.exists(images_folder) else []
                if image_files:
                    image_path = os.path.join(images_folder, random.choice(image_files))
                    try:
                        # Добавляем метку с символом '&' перед изображением
                        label_paragraph = document.add_paragraph("&")
                        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = label_paragraph.runs[0]
                        run.font.size = Pt(base_font_size)
                        run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет шрифта

                        # Устанавливаем свойства форматирования для метки
                        label_paragraph.paragraph_format.keep_with_next = True
                        label_paragraph.paragraph_format.keep_together = True

                        # Добавляем изображение с ограничением по высоте
                        image_paragraph = add_image_to_document(document, image_path, max_height_px=500, base_font_size=base_font_size)


                        # Устанавливаем свойства форматирования для параграфа с изображением
                        image_paragraph.paragraph_format.keep_together = True
                        image_paragraph.paragraph_format.keep_with_next = True

                        # Добавляем подпись к изображению
                        caption_type = random.choice(["Рис.", "Рисунок", "Рис. №", "Рисунок №",])
                        caption_text = f"{caption_type} {random.randint(1, 100)} — {fake.sentence(nb_words=random.randint(3, 7))}"
                        caption_paragraph = document.add_paragraph(caption_text)
                        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption_paragraph.paragraph_format.keep_with_next = True
                        # Устанавливаем размер шрифта для подписи
                        for run in caption_paragraph.runs:
                            run.font.size = Pt(base_font_size)
                    except Exception as e:
                        print(f"Ошибка при добавлении изображения: {e}")

        def add_numbered_list():
            global footnote_num
            numbered_paragraphs = []
            fully_indented = random.choice([True, False])
            # Определяем убирать ли отступы до и после списка
            remove_bef_and_aft_spacing = random.choice([True, False])
            num_of_items = random.randint(3, 7)
            indent = 1
            if use_columns:
                if columns == 2:
                    indent = 0.75
                elif columns == 3:
                    indent = 0.5
            # Определяем будет ли список в целом иметь пункты с большим объемом текста
            bigger_items = random.choice([True, False, False])
            
            for item in range(num_of_items):
                list_item = fake.sentence(nb_words=(random.randint(30, 55) if bigger_items else random.randint(5, 20)))                
                if random.choices([True, False], weights=[0.16, 1 - 0.16])[0]:
                    footnote_text = fake.sentence(nb_words=5)
                    # Если элемент списка большой, то для него возможно добавление сноски где-то в середине
                    if bigger_items and random.choice([True, False]):                  
                        item_len = random.randint(30, 55)
                        first_part_len = random.randint(10, item_len - 10)
                        list_item = fake.sentence(nb_words=first_part_len) + f' [{footnote_num}] '
                        list_item += fake.sentence(nb_words=item_len - first_part_len)                      
                    else:
                        list_item += f' [{footnote_num}]'
                    footnotes.append((footnote_num, footnote_text))
                    footnote_num += 1
                if fully_indented:
                    paragraph = document.add_paragraph(list_item, style='List Number')
                else:
                    paragraph = document.add_paragraph(f'{item + 1}. ' + list_item)
                
                numbered_paragraphs.append(paragraph)
                paragraph_format = paragraph.paragraph_format          

                if item == 0 and remove_bef_and_aft_spacing:
                    paragraph_format.space_before = Pt(0)

                if fully_indented:
                    paragraph_format.left_indent = Cm(1)
                else:
                    paragraph_format.first_line_indent = Cm(indent)               
                    if item != num_of_items - 1:
                        paragraph_format.space_after = Pt(0)

                if item == num_of_items - 1:
                    # Добавляем в конец списка невидимый символ @
                    run = paragraph.add_run('@')
                    run.font.size = Pt(1)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет          
                    if remove_bef_and_aft_spacing:
                        paragraph_format.space_after = Pt(0)
            
                paragraph_format.line_spacing = doc_line_spacing
                run = paragraph.runs[0]
                run.font.size = Pt(base_font_size)
                run.font.name = 'Times New Roman'

            # Устанавливаем размер шрифта для номеров списка
            set_numbering_font_size(document, base_font_size=base_font_size)

        def add_bulleted_list():
            global footnote_num
            bulleted_paragraphs = []
            fully_indented = random.choice([True, False])
            remove_bef_and_aft_spacing = random.choice([True, False])
            num_of_items = random.randint(3, 7)
            indent = 1
            if use_columns:
                if columns == 2:
                    indent = 0.75
                elif columns == 3:
                    indent = 0.5
            # Определяем будет ли список в целом иметь пункты с большим объемом текста
            bigger_items = random.choice([True, False, False])

            for item in range(num_of_items):
                list_item = fake.sentence(nb_words=(random.randint(30, 55) if bigger_items else random.randint(5, 20)))
                if random.choices([True, False], weights=[0.16, 1 - 0.16])[0]:
                    footnote_text = fake.sentence(nb_words=5)
                    # Если элемент списка большой, то для него возможно добавление сноски где-то в середине
                    if bigger_items and random.choice([True, False]):                  
                        item_len = random.randint(30, 55)
                        first_part_len = random.randint(10, item_len - 10)
                        list_item = fake.sentence(nb_words=first_part_len) + f' [{footnote_num}] '
                        list_item += fake.sentence(nb_words=item_len - first_part_len)                      
                    else:
                        list_item += f' [{footnote_num}]'
                    footnotes.append((footnote_num, footnote_text))
                    footnote_num += 1
                if fully_indented:
                    paragraph = document.add_paragraph(list_item, style='List Bullet')
                else:
                    paragraph = document.add_paragraph('• ' + list_item)

                bulleted_paragraphs.append(paragraph)
                paragraph_format = paragraph.paragraph_format

                if item == 0 and remove_bef_and_aft_spacing:
                    paragraph_format.space_before = Pt(0)

                if fully_indented:
                    paragraph_format.left_indent = Cm(1)
                else:
                    paragraph_format.first_line_indent = Cm(indent)               
                    if item != num_of_items - 1:
                        paragraph_format.space_after = Pt(0)

                if item == num_of_items - 1:
                    # Добавляем в конец списка невидимый символ @
                    run = paragraph.add_run('@')
                    run.font.size = Pt(1)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # Белый цвет          
                    if remove_bef_and_aft_spacing:
                        paragraph_format.space_after = Pt(0)

                paragraph_format.line_spacing = doc_line_spacing  # Фиксированное межстрочное расстояние
                run = paragraph.runs[0]
                run.font.size = Pt(base_font_size)
                run.font.name = 'Times New Roman'

            # Устанавливаем размер шрифта для маркеров списка
            set_numbering_font_size(document, base_font_size=base_font_size)

        def add_formula():
            if random.choice([True, False]):
                size_choice = random.choice(["normal", "small", "smallest"])
                num_equations = random.randint(1, 3)  # Случайное число формул (1-3)
                for i in range(num_equations):
                    equation_str = random.choice(FORMULAS)
                    add_equation_to_docx(document, equation_str, caption=True, size_img=size_choice)

        # Список функций для добавления элементов
        elements = [add_text_paragraph, add_numbered_list, add_bulleted_list, add_text_paragraph, add_formula]

        if not use_columns:
            elements.append(add_table)
            elements.append(add_image_or_graph)

        # Перемешиваем элементы
        random.shuffle(elements)
        i = 1
        while (elements[0] ==  add_bulleted_list or elements[0] == add_numbered_list):
            tmp = elements[0]
            elements[0] = elements[i]
            elements[i] = tmp
            i += 1
            # Выполняем функции в случайном порядке
        for element in elements:
            element()

        # # Обновляем footnote_num
        # footnote_num = footnote_num[0]


    # Добавляем сноски в конец документа
    add_footnotes_section(document, footnotes, base_font_size=base_font_size)

    # Сохраняем документ в папку 'docx'
    document.save(f'docx/document_{doc_num}.docx')
    print(f"Документ {doc_num} успешно сгенерирован.")
