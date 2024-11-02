import random
from collections import OrderedDict

# pip install python-docx
# pip install Faker
# pip install mimesis

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from faker import Faker


def change_orientation():
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height

    return current_section


def delete_table(table):
        document.tables[table]._element.getparent().remove(document.tables[table]._element)


def get_table_styles():
    styles = document.styles
    table = document.add_table(rows = 1,cols = 1)
    table_styles = []
    for s in styles:
        try:
            table.style = s
            table_styles.append(s)
        except:
            pass
    delete_table(0)
    return table_styles


COLORS = """000000 000080 00008B 0000CD 0000FF 006400 008000 008080 008B8B 00BFFF 00CED1 
00FA9A 00FF00 00FF7F 00FFFF 00FFFF 191970 1E90FF 20B2AA 228B22 2E8B57 2F4F4F 32CD32 3CB371 
40E0D0 4169E1 4682B4 483D8B 48D1CC 4B0082 556B2F 5F9EA0 6495ED 66CDAA 696969 6A5ACD 6B8E23 
708090 778899 7B68EE 7CFC00 7FFF00 7FFFD4 800000 800080 808000 808080 87CEEB 87CEFA 8A2BE2 
8B0000 8B008B 8B4513 8FBC8F 90EE90 9370D8 9400D3 98FB98 9932CC 9ACD32 A0522D A52A2A A9A9A9 
ADD8E6 ADFF2F AFEEEE B0C4DE B0E0E6 B22222 B8860B BA55D3 BC8F8F BDB76B C0C0C0 C71585 CD5C5C 
CD853F D2691E D2B48C D3D3D3 D87093 D8BFD8 DA70D6 DAA520 DC143C DCDCDC DDA0DD DEB887 E0FFFF 
E6E6FA E9967A EE82EE EEE8AA F08080 F0E68C F0F8FF F0FFF0 F0FFFF F4A460 F5DEB3 F5F5DC F5F5F5 
F5FFFA F8F8FF FA8072 FAEBD7 FAF0E6 FAFAD2 FDF5E6 FF0000 FF00FF FF00FF FF1493 FF4500 FF6347 
FF69B4 FF7F50 FF8C00 FFA07A FFA500 FFB6C1 FFC0CB FFD700 FFDAB9 FFDEAD FFE4B5 FFE4C4 FFE4E1 
FFEBCD FFEFD5 FFF0F5 FFF5EE FFF8DC FFFACD FFFAF0 FFFAFA FFFF00 FFFFE0 FFFFF0 FFFFFF"""


locales = OrderedDict([
    ('en-US', 1),
    ('ru-RU', 2),
])
fake = Faker(locales)

document = Document()
num_tables = 5

colors_list = COLORS.split()
table_styles = get_table_styles()

for i in range(num_tables):
    # для каждой таблицы создаем новую страницу, кроме 1й
    if i != 0:
        document.add_section(WD_SECTION.NEW_PAGE)

    # случайным образом изменяем ориентацию страницы
    if random.randint(0, 1) > 0:
        change_orientation()

    # добавляем сгенерированный текст от 50 до 1000 символов
    document.add_paragraph(fake.text(max_nb_chars = random.randint(50, 1000)))

    # вставляем пустую таблицу от 5 до 10 строк и колонок
    table = document.add_table(
        rows = random.randint(5, 10),
        cols = random.randint(5, 10)
    )

    # случайным образом применяем выравнивание таблицы
    alig = random.randint(0, 2)
    if alig == 0:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
    elif alig == 1:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    else:
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

    # случайным образом изменяем стиль таблицы
    if random.randint(0,1) > 0:
        table.style = table_styles[random.randint(0, len(table_styles) - 1)]

        # заполняем ячейки текстом и числами
        for idx_row, row in enumerate(table.rows):
            for idx_cell, cell in enumerate(row.cells):
                if random.randint(0, 3) > 2:
                    cell.text = fake.text(max_nb_chars = random.randint(5, 50))
                else:
                    cell.text = str(fake.random_number())

    # таблицы с чередующимися колонками в разных стилях и включенной сеткой
    else:
        table.style = 'Table Grid'

        color_row_1 = colors_list[random.randint(0, len(colors_list) - 1)]
        color_row_2 = colors_list[random.randint(0, len(colors_list) - 1)]

        # случайны цвет текста - белый или черный
        if random.randint(0, 1) > 0:
            font_color = RGBColor(255, 255, 255)
        else:
            font_color = RGBColor(0, 0, 0)

        for idx_row, row in enumerate(table.rows):
            if idx_row % 2 == 0:
                color = color_row_1
            else:
                color = color_row_2

            # заполняем ячейки текстом и числами
            for idx_cell, cell in enumerate(row.cells):
                if random.randint(0, 3) > 2:
                    cell.text = fake.text(max_nb_chars = random.randint(5, 50))
                else:
                    cell.text = str(fake.random_number())

                # заливка ячейки цветом
                cell_xml_element = table.rows[idx_row].cells[idx_cell]._tc
                table_cell_properties = cell_xml_element.get_or_add_tcPr()
                shade_obj = OxmlElement('w:shd')
                shade_obj.set(qn('w:fill'), f'{color.lower()}')
                table_cell_properties.append(shade_obj)

                # применение цвета текста
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.color.rgb = font_color

    # добавляем случайный текст в конец документа
    document.add_paragraph(fake.text(max_nb_chars = random.randint(50, 1000)))

# сохраняем документ
document.save(f'demo{i}.docx')
