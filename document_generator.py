import os
import random
from collections import OrderedDict
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from faker import Faker
from mimesis import Person, Datetime, Text, Numeric
from mimesis.enums import Locale

# Создаем директорию 'docx', если она не существует
if not os.path.exists('docx'):
    os.makedirs('docx')

# Список цветов для заливки ячеек таблиц
COLORS = """000000 000080 00008B 0000CD 0000FF 006400 008000 008080 008B8B 00BFFF 00CED1 
00FA9A 00FF00 00FF7F 00FFFF 191970 1E90FF 20B2AA 228B22 2E8B57 2F4F4F 32CD32 3CB371 
40E0D0 4169E1 4682B4 483D8B 48D1CC 4B0082 556B2F 5F9EA0 6495ED 66CDAA 696969 6A5ACD 6B8E23 
708090 778899 7B68EE 7CFC00 7FFF00 7FFFD4 800000 800080 808000 808080 87CEEB 87CEFA 8A2BE2 
8B0000 8B008B 8B4513 8FBC8F 90EE90 9370D8 9400D3 98FB98 9932CC 9ACD32 A0522D A52A2A A9A9A9 
ADD8E6 ADFF2F AFEEEE B0C4DE B0E0E6 B22222 B8860B BA55D3 BC8F8F BDB76B C0C0C0 C71585 CD5C5C 
CD853F D2691E D2B48C D3D3D3 D87093 D8BFD8 DA70D6 DAA520 DC143C DCDCDC DDA0DD DEB887 E0FFFF 
E6E6FA E9967A EE82EE EEE8AA F08080 F0E68C F0F8FF F0FFF0 F0FFFF F4A460 F5DEB3 F5F5DC F5F5F5 
F5FFFA F8F8FF FA8072 FAEBD7 FAF0E6 FAFAD2 FDF5E6 FF0000 FF00FF FF1493 FF4500 FF6347 
FF69B4 FF7F50 FF8C00 FFA07A FFA500 FFB6C1 FFC0CB FFD700 FFDAB9 FFDEAD FFE4B5 FFE4C4 FFE4E1 
FFEB7D FFEBCD FFEFD5 FFF0F5 FFF5EE FFF8DC FFFACD FFFAF0 FFFAFA FFFF00 FFFFE0 FFFFF0 FFFFFF"""

colors_list = COLORS.split()

# Инициализируем Faker и Mimesis
locales = OrderedDict([
    ('en-US', 1),
    ('ru-RU', 2),
])
fake = Faker(locales)
person = Person(Locale.RU)
datetime_mimesis = Datetime()
text_mimesis = Text(Locale.RU)
numeric = Numeric()

def get_table_styles(document):
    """
    Получает список доступных стилей таблиц в документе.
    """
    styles = document.styles
    table = document.add_table(rows=1, cols=1)
    table_styles = []
    for s in styles:
        try:
            table.style = s
            table_styles.append(s.name)
        except:
            pass
    # Удаляем временную таблицу
    table._element.getparent().remove(table._element)
    return table_styles

def change_orientation(section):
    """
    Меняет ориентацию страницы на альбомную.
    """
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

def add_footnote(paragraph, footnote_text, footnote_num, footnotes):
    """
    Добавляет сноску в абзац и сохраняет её текст.
    """
    footnote_mark = paragraph.add_run(f'[{footnote_num}]')
    footnote_mark.font.superscript = True
    footnotes.append((footnote_num, footnote_text))

def add_headers_footers(document):
    """
    Добавляет колонтитулы (верхний и нижний) в документ.
    """
    for section in document.sections:
        header = section.header
        footer = section.footer
        header_text = fake.sentence(nb_words=5)
        footer_text = fake.sentence(nb_words=5)
        header_para = header.paragraphs[0]
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para = footer.paragraphs[0]
        footer_para.text = footer_text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_footnotes_section(document, footnotes):
    """
    Добавляет раздел с примечаниями (сносками) в конец документа.
    """
    if footnotes:
        document.add_page_break()
        footnote_heading = document.add_paragraph('Примечания')
        footnote_heading.style = 'Heading 1'
        for num, text in footnotes:
            footnote_paragraph = document.add_paragraph()
            footnote_run = footnote_paragraph.add_run(f'[{num}] {text}')
            footnote_run.font.size = Pt(8)

def generate_document(doc_num):
    try:
        print(f"Генерация документа {doc_num}...")
        document = Document()

        footnotes = []

        # Получаем стили таблиц
        table_styles = get_table_styles(document)

        # Случайно устанавливаем ориентацию страницы
        section = document.sections[-1]
        if random.choice([True, False]):
            change_orientation(section)

        # Случайно устанавливаем количество колонок (1 или 2-3)
        if random.choice([True, False]):
            sectPr = section._sectPr
            cols_elem = sectPr.xpath('./w:cols')
            if cols_elem:
                cols = cols_elem[0]
            else:
                cols = OxmlElement('w:cols')
                sectPr.append(cols)
            num_columns = random.randint(2, 3)
            cols.set(qn('w:num'), str(num_columns))

        # Определяем порядок добавления элементов
        elements = ['heading', 'text', 'table', 'list', 'image', 'formula']
        random.shuffle(elements)  # Перемешиваем порядок

        for element in elements:
            if element == 'heading':
                # Добавляем заголовок (различных уровней)
                heading_level = random.randint(1, 3)  # Heading 1, 2 или 3
                heading_text = fake.sentence(nb_words=random.randint(3, 7))
                heading_paragraph = document.add_heading(heading_text, level=heading_level)
                heading_paragraph.alignment = random.choice([WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
                
                # Применяем стиль к заголовку
                heading_style = heading_paragraph.style
                heading_style.font.size = Pt(random.randint(16, 24))
                heading_style.font.bold = random.choice([True, False])
                heading_style.font.italic = random.choice([True, False])
                heading_style.font.name = 'Times New Roman'

            elif element == 'text':
                # Добавляем несколько абзацев текста
                for _ in range(random.randint(1, 3)):
                    text_length = random.randint(100, 500)
                    text_content = fake.text(max_nb_chars=text_length)
                    text_paragraph = document.add_paragraph()
                    text_paragraph.alignment = random.choice([WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT])
                    text_paragraph.paragraph_format.first_line_indent = Pt(20) if random.choice([True, False]) else Pt(0)
                    
                    # Разбиваем текст на предложения
                    sentences = text_content.split('. ')
                    footnote_num = len(footnotes) + 1  # Продолжаем нумерацию с существующих сносок
                    for sentence in sentences:
                        if sentence.strip() == '':
                            continue
                        if random.choice([True, False, False]):  # Увеличиваем вероятность добавления сносок
                            # Добавляем сноску
                            footnote_text = fake.sentence(nb_words=5)
                            text_paragraph.add_run(sentence + '. ')
                            add_footnote(text_paragraph, footnote_text, footnote_num, footnotes)
                            footnote_num +=1
                        else:
                            text_paragraph.add_run(sentence + '. ')

            elif element == 'table':
                # Добавляем подпись к таблице
                table_caption_text = f"Таблица {random.randint(1,100)}. {fake.sentence(nb_words=5)}"
                caption_style = 'CaptionStyle'  # Используем предопределённый стиль
                if caption_style not in [s.name for s in document.styles]:
                    # Создаём стиль, если его нет
                    try:
                        caption_style_obj = document.styles.add_style(caption_style, WD_STYLE_TYPE.PARAGRAPH)
                        caption_style_obj.font.size = Pt(10)
                        caption_style_obj.font.italic = True
                        caption_style_obj.font.name = 'Times New Roman'
                    except Exception as e:
                        print(f"Ошибка при создании стиля {caption_style}: {e}")
                
                caption_paragraph = document.add_paragraph(table_caption_text, style=caption_style)
                caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Добавляем таблицу
                rows = random.randint(2, 10)
                cols = random.randint(2, 5)
                table = document.add_table(rows=rows, cols=cols)
            
                # Случайно выбираем стиль таблицы
                if table_styles:
                    try:
                        table.style = random.choice(table_styles)
                    except IndexError:
                        table.style = 'Table Grid'
                else:
                    try:
                        table.style = 'Table Grid'
                    except KeyError:
                        pass  # Если 'Table Grid' недоступен

                table.alignment = random.choice([WD_TABLE_ALIGNMENT.LEFT, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.RIGHT])
            
                # Заполняем таблицу данными
                for idx_row, row in enumerate(table.rows):
                    for idx_cell, cell in enumerate(row.cells):
                        if random.randint(0, 3) > 2:
                            cell.text = fake.text(max_nb_chars=random.randint(5, 50))
                        else:
                            cell.text = str(fake.random_number(digits=5))
                    
                        # Случайно заливаем ячейки цветом
                        if random.choice([True, False]):
                            color = random.choice(colors_list)
                            cell_xml_element = cell._tc
                            table_cell_properties = cell_xml_element.get_or_add_tcPr()
                            shade_obj = OxmlElement('w:shd')
                            shade_obj.set(qn('w:fill'), f'{color.lower()}')
                            table_cell_properties.append(shade_obj)
        
                        # Применяем случайный цвет текста
                        if random.choice([True, False]):
                            font_color = random.choice([RGBColor(255, 255, 255), RGBColor(0, 0, 0)])
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = font_color

            elif element == 'list':
                # Добавляем нумерованный список
                for _ in range(random.randint(1, 3)):
                    list_item = fake.sentence(nb_words=6)
                    p = document.add_paragraph(list_item, style='List Number')
                    p.paragraph_format.left_indent = Pt(20)
                    p.paragraph_format.space_after = Pt(0)
            
                # Добавляем маркированный список
                for _ in range(random.randint(1, 3)):
                    list_item = fake.sentence(nb_words=6)
                    p = document.add_paragraph(list_item, style='List Bullet')
                    p.paragraph_format.left_indent = Pt(20)
                    p.paragraph_format.space_after = Pt(0)

            elif element == 'image':
                # Добавляем изображение с подписью
                # Убедитесь, что у вас есть изображения в папке 'images'
                image_files = os.listdir('images') if os.path.exists('images') else []
                if image_files:
                    image_path = os.path.join('images', random.choice(image_files))
                    try:
                        document.add_picture(image_path, width=Inches(5))
                        image_caption_text = f"Рисунок {random.randint(1,100)} - {fake.sentence(nb_words=5)}"
                        image_caption_paragraph = document.add_paragraph(image_caption_text)
                        image_caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Ошибка при добавлении изображения: {e}")

            elif element == 'formula':
                # Добавляем формулу (как изображение)
                # Убедитесь, что у вас есть изображения формул в папке 'formulas'
                formula_files = os.listdir('formulas') if os.path.exists('formulas') else []
                if formula_files:
                    formula_path = os.path.join('formulas', random.choice(formula_files))
                    try:
                        document.add_picture(formula_path, width=Inches(5))
                        formula_caption_text = f"Формула {random.randint(1,100)}"
                        formula_caption_paragraph = document.add_paragraph(formula_caption_text)
                        formula_caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"Ошибка при добавлении формулы: {e}")

        # Добавляем колонтитулы
        add_headers_footers(document)

        # Добавляем сноски в конец документа
        add_footnotes_section(document, footnotes)

        # Добавляем случайный текст в конец документа
        document.add_paragraph(fake.text(max_nb_chars=random.randint(50, 1000)))

        # Сохраняем документ
        document.save(f'docx/document_{doc_num}.docx')
        print(f"Документ {doc_num} успешно сгенерирован.")
    
    except Exception as e:
        print(f"Ошибка при генерации документа {doc_num}: {e}")

# Генерируем документы (для примера сгенерируем 3 документа)
for i in range(1, 11):
    generate_document(i)
